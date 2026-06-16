import io
import pdfplumber
import fitz
from docx import Document


def parse_pdf(content: bytes) -> str:
    text = _parse_pdf_pdfplumber(content)
    if not text or len(text.strip()) < 50:
        text = _parse_pdf_pymupdf(content)
    return clean_text(text)


def _parse_pdf_pdfplumber(content: bytes) -> str:
    parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
    return "\n".join(parts)


def _parse_pdf_pymupdf(content: bytes) -> str:
    parts = []
    doc = fitz.open(stream=content, filetype="pdf")
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts)


def parse_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return clean_text("\n".join(parts))


def clean_text(text: str) -> str:
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            cleaned.append(stripped)
    return "\n".join(cleaned)
